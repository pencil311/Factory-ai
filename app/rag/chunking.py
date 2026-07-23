"""Structure-aware chunking.

Naive fixed-size chunking destroys exactly the content that matters most in
maintenance manuals. Two failure modes drive this design:

* **Tables get shredded.** An error-code table split down the middle leaves
  "E104" in one chunk and "Spindle over-temperature — check coolant flow" in
  another. Neither passage answers the question, and the retriever cannot
  reassemble them. Tables are therefore detected and emitted whole, even when
  that exceeds the token target.
* **Headings get orphaned.** A passage reading "Torque to 45 Nm" is useless
  without knowing it sits under "6.2 Drive Roller Bearing Replacement". The
  nearest heading is carried into every chunk as ``section_title``.

Sentences are never split: the splitter accumulates whole sentences and stops
before overflowing, so a chunk boundary never lands mid-clause.

Token counting is a deliberate approximation (~4 chars/token) rather than a
real tokenizer — it keeps this module dependency-free and the target is a soft
one. Chunks are sized for retrieval quality, not for a hard model limit.
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Iterable, Optional, Sequence

DEFAULT_TARGET_TOKENS = 500
DEFAULT_OVERLAP_TOKENS = 80
#: Characters per token, averaged over technical English. Good enough for
#: sizing; nothing downstream depends on it being exact.
CHARS_PER_TOKEN = 4

#: A table row needs at least this many column separators to count.
_MIN_TABLE_COLUMNS = 2
#: Consecutive delimited lines needed before we treat a run as a table.
_MIN_TABLE_ROWS = 2

#: Markdown ATX headings, numbered headings ("6.2 Bearing Replacement"), and
#: ALL-CAPS standalone lines, which manuals use constantly.
_HEADING_PATTERNS = (
    re.compile(r"^\s{0,3}(#{1,6})\s+(?P<title>\S.*?)\s*#*\s*$"),
    re.compile(r"^\s*(?P<num>\d+(?:\.\d+)*)[.)]?\s+(?P<title>[A-Z][^\n]{2,80})$"),
    re.compile(r"^\s*(?P<title>[A-Z0-9][A-Z0-9 \-/&(),.]{3,70})\s*$"),
)

_SENTENCE_END = re.compile(r"(?<=[.!?])\s+(?=[A-Z(\[])|(?<=[.!?])\s*\n")


def estimate_tokens(text: str) -> int:
    """Approximate token count for sizing decisions."""
    return max(1, len(text) // CHARS_PER_TOKEN)


@dataclass
class ChunkDraft:
    """A chunk before it is embedded and persisted."""

    text: str
    page_number: Optional[int] = None
    section_title: Optional[str] = None
    chunk_index: int = 0
    is_table: bool = False
    token_count: int = 0

    def __post_init__(self) -> None:
        if not self.token_count:
            self.token_count = estimate_tokens(self.text)


@dataclass
class ParsedPage:
    """One page of a parsed source document."""

    page_number: int
    text: str


@dataclass
class ParsedDocument:
    """A source file reduced to pages of text."""

    pages: list[ParsedPage] = field(default_factory=list)

    @property
    def page_count(self) -> int:
        return len(self.pages)

    @property
    def full_text(self) -> str:
        return "\n\n".join(p.text for p in self.pages)


# ---------------------------------------------------------------------------
# Structure detection
# ---------------------------------------------------------------------------
def detect_heading(line: str) -> Optional[str]:
    """Return the heading title if ``line`` looks like a heading.

    Ordered most-specific first: an ATX heading is unambiguous, a numbered
    heading is near-certain, and the ALL-CAPS rule is the loosest so it is
    tried last and guarded against matching table rows or sentences.
    """
    stripped = line.strip()
    if not stripped or len(stripped) > 90:
        return None
    if is_table_row(stripped):
        return None

    for pattern in _HEADING_PATTERNS:
        match = pattern.match(line)
        if not match:
            continue
        title = match.group("title").strip()
        # Trailing sentence or continuation punctuation means prose, not a
        # heading. Numbered *list items* ("4. Stored energy — hydraulic
        # pressure, compressed air,") match the numbered-heading shape but
        # run long and break across lines; a real heading is short and ends
        # cleanly. Without this, list items become section titles and every
        # chunk under them is mis-cited.
        if title[-1:] in ".,;:":
            return None
        if match.groupdict().get("num") and len(title) > 60:
            return None
        # Keep the section number: "6.2 Bearing Replacement" is a far more
        # useful citation than "Bearing Replacement" alone, and it matches how
        # an ATX heading ("## 6.2 Bearing Replacement") already reads.
        groups = match.groupdict()
        number = (groups.get("num") or "").strip()
        if number and not title.startswith(number):
            return f"{number} {title}"
        return title
    return None


def is_table_row(line: str) -> bool:
    """True when a line looks like a row of a delimited table.

    Recognises markdown pipes, and the multi-space column alignment that PDF
    text extraction produces from ruled tables.
    """
    stripped = line.strip()
    if not stripped:
        return False
    if stripped.count("|") >= _MIN_TABLE_COLUMNS:
        return True
    # A markdown separator row: |---|---| or ----  ----
    if re.fullmatch(r"[\s|+:-]{4,}", stripped) and "-" in stripped:
        return True
    # Two or more columns separated by runs of whitespace, e.g. extracted PDF.
    if len(re.findall(r"\S\s{2,}\S", stripped)) >= _MIN_TABLE_COLUMNS - 1:
        return len(stripped.split()) >= 3
    return False


def _split_sentences(text: str) -> list[str]:
    """Split into sentences, keeping their trailing whitespace semantics.

    Falls back to line-splitting for text with no sentence punctuation, which
    is common in bulleted procedures.
    """
    parts = [p.strip() for p in _SENTENCE_END.split(text) if p and p.strip()]
    return parts or [text.strip()]


@dataclass
class _Block:
    """A run of lines that is either prose or a table."""

    lines: list[str]
    is_table: bool
    heading: Optional[str]
    page_number: int

    @property
    def text(self) -> str:
        return "\n".join(self.lines).strip()


def _blocks_for_page(
    page: ParsedPage, inherited_heading: Optional[str]
) -> tuple[list[_Block], Optional[str]]:
    """Split one page into prose/table blocks, tracking the current heading."""
    blocks: list[_Block] = []
    current_heading = inherited_heading
    buffer: list[str] = []
    buffer_is_table = False

    def flush() -> None:
        nonlocal buffer, buffer_is_table
        if buffer and any(line.strip() for line in buffer):
            blocks.append(
                _Block(
                    lines=list(buffer),
                    is_table=buffer_is_table,
                    heading=current_heading,
                    page_number=page.page_number,
                )
            )
        buffer = []
        buffer_is_table = False

    lines = page.text.splitlines()
    index = 0
    while index < len(lines):
        line = lines[index]

        heading = detect_heading(line)
        if heading is not None:
            flush()
            current_heading = heading
            index += 1
            continue

        if is_table_row(line):
            # Look ahead: a lone delimited line is probably not a table.
            run_end = index
            while run_end < len(lines) and is_table_row(lines[run_end]):
                run_end += 1
            if run_end - index >= _MIN_TABLE_ROWS:
                flush()
                blocks.append(
                    _Block(
                        lines=lines[index:run_end],
                        is_table=True,
                        heading=current_heading,
                        page_number=page.page_number,
                    )
                )
                index = run_end
                continue

        buffer.append(line)
        index += 1

    flush()
    return blocks, current_heading


# ---------------------------------------------------------------------------
# Chunking
# ---------------------------------------------------------------------------
def chunk_document(
    parsed: ParsedDocument,
    target_tokens: int = DEFAULT_TARGET_TOKENS,
    overlap_tokens: int = DEFAULT_OVERLAP_TOKENS,
) -> list[ChunkDraft]:
    """Chunk a parsed document, preserving tables and section context."""
    drafts: list[ChunkDraft] = []
    heading: Optional[str] = None

    for page in parsed.pages:
        blocks, heading = _blocks_for_page(page, heading)
        for block in blocks:
            if block.is_table:
                # Whole and unsplit, whatever its size. A truncated error-code
                # table is worse than a long chunk.
                drafts.append(
                    ChunkDraft(
                        text=block.text,
                        page_number=block.page_number,
                        section_title=block.heading,
                        is_table=True,
                    )
                )
                continue
            drafts.extend(
                _chunk_prose(
                    block.text,
                    page_number=block.page_number,
                    section_title=block.heading,
                    target_tokens=target_tokens,
                    overlap_tokens=overlap_tokens,
                )
            )

    for position, draft in enumerate(drafts):
        draft.chunk_index = position
    return drafts


def _chunk_prose(
    text: str,
    page_number: int,
    section_title: Optional[str],
    target_tokens: int,
    overlap_tokens: int,
) -> list[ChunkDraft]:
    """Accumulate whole sentences up to the target, with sentence overlap."""
    text = text.strip()
    if not text:
        return []

    sentences = _split_sentences(text)
    chunks: list[ChunkDraft] = []
    current: list[str] = []
    current_tokens = 0

    def emit() -> None:
        if not current:
            return
        chunks.append(
            ChunkDraft(
                text=" ".join(current).strip(),
                page_number=page_number,
                section_title=section_title,
            )
        )

    for sentence in sentences:
        sentence_tokens = estimate_tokens(sentence)

        if current and current_tokens + sentence_tokens > target_tokens:
            emit()
            # Carry a tail of whole sentences forward so a passage spanning a
            # boundary is still retrievable from either side.
            carried: list[str] = []
            carried_tokens = 0
            for previous in reversed(current):
                previous_tokens = estimate_tokens(previous)
                if carried_tokens + previous_tokens > overlap_tokens:
                    break
                carried.insert(0, previous)
                carried_tokens += previous_tokens
            current = carried
            current_tokens = carried_tokens

        current.append(sentence)
        current_tokens += sentence_tokens

    emit()
    return chunks


# ---------------------------------------------------------------------------
# Parsing
# ---------------------------------------------------------------------------
SUPPORTED_SUFFIXES = (".pdf", ".txt", ".md", ".markdown", ".docx")


class UnsupportedFileType(ValueError):
    """Raised for a file extension the parser does not handle."""


def parse_bytes(data: bytes, filename: str) -> ParsedDocument:
    """Parse raw file bytes into pages of text, dispatching on extension."""
    suffix = Path(filename).suffix.lower()
    if suffix == ".pdf":
        return _parse_pdf(data)
    if suffix == ".docx":
        return _parse_docx(data)
    if suffix in (".txt", ".md", ".markdown"):
        return _parse_text(data.decode("utf-8", errors="replace"))
    raise UnsupportedFileType(
        f"Cannot parse '{filename}'. Supported types: {', '.join(SUPPORTED_SUFFIXES)}"
    )


def _parse_text(text: str) -> ParsedDocument:
    """Treat form feeds as page breaks; otherwise one page."""
    pages = text.split("\f") if "\f" in text else [text]
    return ParsedDocument(
        pages=[
            ParsedPage(page_number=i, text=page)
            for i, page in enumerate(pages, start=1)
            if page.strip()
        ]
    )


def _parse_pdf(data: bytes) -> ParsedDocument:
    """Extract text per page with pypdf, preserving page numbers for citation."""
    import io

    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    pages = []
    for number, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        if text.strip():
            pages.append(ParsedPage(page_number=number, text=text))
    return ParsedDocument(pages=pages)


def _parse_docx(data: bytes) -> ParsedDocument:
    """Extract paragraphs and tables from a .docx.

    Word tables are rendered back to pipe-delimited rows so the table detector
    keeps them whole, rather than losing their structure on the way in.
    """
    import io

    import docx  # python-docx

    document = docx.Document(io.BytesIO(data))
    lines: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style = (paragraph.style.name or "").lower()
        # Re-emit Word headings as markdown so detect_heading sees them.
        if style.startswith("heading") or style == "title":
            lines.append(f"## {text}")
        else:
            lines.append(text)

    for table in document.tables:
        lines.append("")
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            lines.append(" | ".join(cells))
        lines.append("")

    return ParsedDocument(pages=[ParsedPage(page_number=1, text="\n".join(lines))])


def parse_and_chunk(
    data: bytes,
    filename: str,
    target_tokens: int = DEFAULT_TARGET_TOKENS,
    overlap_tokens: int = DEFAULT_OVERLAP_TOKENS,
) -> tuple[ParsedDocument, list[ChunkDraft]]:
    """Parse then chunk, returning both so the caller can record page_count."""
    parsed = parse_bytes(data, filename)
    return parsed, chunk_document(parsed, target_tokens, overlap_tokens)
